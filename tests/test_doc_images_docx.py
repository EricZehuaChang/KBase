"""多模态扩展：docx 插图 caption 级锚定（命中章节才附图）、图片文件文档
命中即附原图（OCR 路径）。"""
import io
import json as _json

import pytest
from docx import Document as DocxDocument
from docx.shared import Inches
from fastapi.testclient import TestClient
from PIL import Image

from kbase.api.main import create_app
from tests.test_api import CFG, FakeLLM
from tests.test_ocr_routing import FakeOCR


def _jpeg(w=320, h=240) -> bytes:
    buf = io.BytesIO()
    Image.frombytes(
        "RGB", (w, h),
        bytes((i * 37 + j * 11) % 256 for j in range(h) for i in range(w * 3))
    ).save(buf, format="JPEG", quality=95)
    return buf.getvalue()


def _client(tmp_path, fake_embedder, ocr_backend=None):
    cfg = tmp_path / "kbase.yaml"
    cfg.write_text(CFG.format(data_dir=str(tmp_path / "data").replace("\\", "/")),
                   encoding="utf-8")
    app = create_app(config_path=cfg, embedder=fake_embedder,
                     llms={"fake": FakeLLM()}, reranker=False, auth="off",
                     ocr_backend=ocr_backend)
    return TestClient(app)


def _citations_for(c, kb, question):
    citations = []
    with c.stream("POST", f"/api/kb/{kb}/query",
                  json={"question": question}) as resp:
        event = ""
        for line in resp.iter_lines():
            if line.startswith("event:"):
                event = line[6:].strip()
            elif line.startswith("data:") and event == "citations":
                citations = _json.loads(line[5:].strip())
                break
    return citations


def test_docx_images_anchored_to_heading(tmp_path, fake_embedder):
    """插图锚到所在章节：命中'报销流程'章节的引用附图，命中'考勤'章节的不附。"""
    doc = DocxDocument()
    doc.add_heading("考勤制度", level=1)
    doc.add_paragraph("迟到三次记警告一次，情节严重记旷工。")
    doc.add_heading("报销流程", level=1)
    doc.add_paragraph("差旅报销先提交行程单，再走OA两级审批，附发票原件。")
    doc.add_picture(io.BytesIO(_jpeg()), width=Inches(3))   # 流程图在报销章节下
    buf = io.BytesIO()
    doc.save(buf)

    c = _client(tmp_path, fake_embedder)
    kb = c.post("/api/kb", json={"name": "库"}).json()["id"]
    r = c.post(f"/api/kb/{kb}/documents", files=[
        ("files", ("制度.docx", buf.getvalue(),
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))])
    assert r.status_code == 200
    docs = c.get(f"/api/kb/{kb}/documents").json()
    assert docs[0]["status"] == "ready", docs[0]

    # 命中"报销流程"章节 → 附该章节插图
    hits = _citations_for(c, kb, "差旅报销怎么走流程")
    reimburse = [ci for ci in hits if "报销" in ci["heading_path"]]
    assert reimburse and reimburse[0].get("images"), \
        f"报销章节引用应附图: {hits}"
    img = reimburse[0]["images"][0]
    assert img["width"] == 320

    # 图片直链可取回
    got = c.get(img["url"])
    assert got.status_code == 200 and got.headers["content-type"].startswith("image/")

    # 命中"考勤制度"章节 → 不附图（caption 级锚定，不做全文档洪泛）
    hits2 = _citations_for(c, kb, "迟到三次怎么处理")
    kaoqin = [ci for ci in hits2 if "考勤" in ci["heading_path"]]
    assert kaoqin and not kaoqin[0].get("images")


def test_image_file_doc_attaches_original(tmp_path, fake_embedder):
    """图片文件文档（扫描照上传走 OCR）：命中即附原图本身。"""
    c = _client(tmp_path, fake_embedder, ocr_backend=FakeOCR())
    kb = c.post("/api/kb", json={"name": "库"}).json()["id"]
    r = c.post(f"/api/kb/{kb}/documents",
               files=[("files", ("验收单.png", _jpeg(300, 200), "image/png"))])
    assert r.status_code == 200
    assert c.get(f"/api/kb/{kb}/documents").json()[0]["status"] == "ready"

    citations = _citations_for(c, kb, "扫描件识别出的内容")
    assert citations
    images = citations[0].get("images")
    assert images and images[0]["url"].endswith("/original?disposition=inline")
    assert images[0]["name"] == "验收单.png"
