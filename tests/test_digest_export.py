from docx import Document as DocxDocument

from kbase.db import make_session_factory
from kbase.jobs.digest import build_digest_steps
from kbase.jobs.export_docx import markdown_to_docx
from kbase.jobs.runner import run_job
from kbase.jobs.store import create_job, get_job
from kbase.models import Document


def _sf(tmp_path):
    return make_session_factory(f"sqlite:///{tmp_path}/kb.sqlite")


def _seed_doc(sf, doc_id: str, kb_id: str, filename: str, status: str = "ready"):
    with sf() as s:
        s.add(Document(id=doc_id, kb_id=kb_id, filename=filename,
                       content_hash=f"hash-{doc_id}", status=status))
        s.commit()


def _write_content(files_dir, doc_id: str, text: str):
    out_dir = files_dir / doc_id
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "content.md").write_text(text, encoding="utf-8")


class FakeLLM:
    """按调用顺序返回预置文本；记录收到的 messages 供断言 prompt 内容。"""

    def __init__(self, outputs):
        self.outputs = list(outputs)
        self.calls = []

    async def complete(self, messages, **params):
        self.calls.append(messages)
        return self.outputs[len(self.calls) - 1]


# ---- digest 全流程 ----

def test_digest_full_flow_produces_artifact_with_overview_and_sections(tmp_path):
    sf = _sf(tmp_path)
    _seed_doc(sf, "d1", "kb1", "政策一.docx")
    _seed_doc(sf, "d2", "kb1", "政策二.docx")
    files_dir = tmp_path / "files"
    _write_content(files_dir, "d1", "政策一正文内容，关于住房补贴的规定。")
    _write_content(files_dir, "d2", "政策二正文内容，关于差旅补助的规定。")

    job = create_job(sf, kb_id="kb1", type="digest", params={}, provider=None)
    jobs_dir = tmp_path / "jobs"

    # 两份摘要 + 一段总览
    llm = FakeLLM(["政策一摘要：住房补贴。", "政策二摘要：差旅补助。", "总览：两项政策均涉及员工福利。"])

    steps = build_digest_steps(sf, llm, kb_id="kb1", doc_ids=None, job_id=job["id"],
                               files_dir=files_dir, jobs_dir=jobs_dir, kb_name="彩排知识库")
    run_job(sf, job["id"], steps)

    got = get_job(sf, job["id"])
    assert got["status"] == "done"
    assert got["artifact_path"]

    md = open(got["artifact_path"], encoding="utf-8").read()
    assert "# 彩排知识库文档汇编" in md
    assert "总览" in md
    assert "总览：两项政策均涉及员工福利。" in md
    assert "## 政策一.docx" in md
    assert "## 政策二.docx" in md
    assert "政策一摘要：住房补贴。" in md
    assert "政策二摘要：差旅补助。" in md


def test_digest_missing_content_md_step_fails_others_still_summarized(tmp_path):
    sf = _sf(tmp_path)
    _seed_doc(sf, "d1", "kb1", "有正文.docx")
    _seed_doc(sf, "d2", "kb1", "无正文.docx")
    files_dir = tmp_path / "files"
    _write_content(files_dir, "d1", "有正文的文档内容。")
    # d2 故意不写 content.md

    job = create_job(sf, kb_id="kb1", type="digest", params={}, provider=None)
    jobs_dir = tmp_path / "jobs"

    llm = FakeLLM(["有正文摘要。", "总览：仅一份文档有效摘要。"])

    steps = build_digest_steps(sf, llm, kb_id="kb1", doc_ids=None, job_id=job["id"],
                               files_dir=files_dir, jobs_dir=jobs_dir, kb_name="彩排知识库")
    run_job(sf, job["id"], steps)

    got = get_job(sf, job["id"])
    assert got["status"] == "done_with_errors"
    assert got["artifact_path"]

    steps_state = got["progress"]["steps"]
    statuses = {s["name"]: s["status"] for s in steps_state}
    assert statuses["摘要：有正文.docx"] == "done"
    assert statuses["摘要：无正文.docx"] == "failed"

    md = open(got["artifact_path"], encoding="utf-8").read()
    assert "## 有正文.docx" in md
    assert "有正文摘要。" in md


def test_digest_with_explicit_doc_ids_only_summarizes_selected(tmp_path):
    sf = _sf(tmp_path)
    _seed_doc(sf, "d1", "kb1", "甲.docx")
    _seed_doc(sf, "d2", "kb1", "乙.docx")
    files_dir = tmp_path / "files"
    _write_content(files_dir, "d1", "甲文档内容。")
    _write_content(files_dir, "d2", "乙文档内容。")

    job = create_job(sf, kb_id="kb1", type="digest", params={}, provider=None)
    jobs_dir = tmp_path / "jobs"

    llm = FakeLLM(["甲摘要。", "总览：仅甲。"])

    steps = build_digest_steps(sf, llm, kb_id="kb1", doc_ids=["d1"], job_id=job["id"],
                               files_dir=files_dir, jobs_dir=jobs_dir, kb_name="库")
    run_job(sf, job["id"], steps)

    got = get_job(sf, job["id"])
    assert got["status"] == "done"
    md = open(got["artifact_path"], encoding="utf-8").read()
    assert "## 甲.docx" in md
    assert "## 乙.docx" not in md


def test_digest_summary_prompt_content_length_capped(tmp_path):
    """内容超过 6000 字时，摘要 prompt 中的正文节选只包含前 6000 字（避免超长 prompt）。"""
    sf = _sf(tmp_path)
    _seed_doc(sf, "d1", "kb1", "长文档二.docx")
    files_dir = tmp_path / "files"
    long_text = "内容" * 5000   # 10000 字符
    _write_content(files_dir, "d1", long_text)

    job = create_job(sf, kb_id="kb1", type="digest", params={}, provider=None)
    jobs_dir = tmp_path / "jobs"
    llm = FakeLLM(["摘要。", "总览。"])

    steps = build_digest_steps(sf, llm, kb_id="kb1", doc_ids=None, job_id=job["id"],
                               files_dir=files_dir, jobs_dir=jobs_dir, kb_name="库")
    run_job(sf, job["id"], steps)

    user_message = llm.calls[0][1]["content"]
    # 正文中截取的内容不应超过 6000 字符（原文 10000 字符已被截断）
    content_in_prompt = user_message.split("正文（节选）：\n", 1)[1].rsplit("\n\n请输出", 1)[0]
    assert len(content_in_prompt) == 6000


# ---- markdown_to_docx ----

SAMPLE_MD = """# 标题一

## 标题二

### 标题三

这是一个普通段落，包含**加粗文本**在内。

- 列表项一
- 列表项二
- 列表项三

另一个普通段落。
"""


def test_markdown_to_docx_converts_headings_paragraphs_bullets_bold(tmp_path):
    out_path = tmp_path / "artifact.docx"
    markdown_to_docx(SAMPLE_MD, out_path)

    assert out_path.exists()
    doc = DocxDocument(str(out_path))

    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "Heading 3" in styles

    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullet_paras) == 3
    assert [p.text for p in bullet_paras] == ["列表项一", "列表项二", "列表项三"]

    h1 = next(p for p in doc.paragraphs if p.style.name == "Heading 1")
    assert h1.text == "标题一"
    h2 = next(p for p in doc.paragraphs if p.style.name == "Heading 2")
    assert h2.text == "标题二"
    h3 = next(p for p in doc.paragraphs if p.style.name == "Heading 3")
    assert h3.text == "标题三"

    # 粗体 run 存在
    bold_para = next(p for p in doc.paragraphs if "加粗文本" in p.text)
    assert any(run.bold for run in bold_para.runs)

    # 段落数量合理：2 个普通段落 + 3 个标题 + 3 个列表项 = 8
    assert len(doc.paragraphs) == 8
