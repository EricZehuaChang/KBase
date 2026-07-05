"""Markdown → docx 最小转换（F4）。仅支持产物 Markdown 实际用到的子集：
#/##/### 标题、空行分段的普通段落、`- ` 无序列表、`**bold**` 行内加粗。
其余一律原样落为普通段落——不追求通用 Markdown 兼容，够用即可（YAGNI）。
"""
import re

from docx import Document

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_runs_with_bold(paragraph, text: str) -> None:
    """把 text 中 **bold** 片段拆成加粗 run，其余原样 run。"""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        bold_run = paragraph.add_run(m.group(1))
        bold_run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(md: str, out_path) -> None:
    """将 md 转换为 docx 并写入 out_path。逐行解析：
    - `# ` / `## ` / `### ` → Heading 1/2/3
    - `- ` 开头 → "List Bullet" 样式段落
    - 空行 → 段落分隔（忽略，不产生空段落）
    - 其余非空行 → 普通段落，支持 **bold** 行内加粗
    """
    doc = Document()

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line.strip())

    doc.save(str(out_path))
